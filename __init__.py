import os
import sys
import docx
import datetime
import subprocess
from docx.dml.color import ColorFormat
from docx.text.run import Font, Run
from docx.shared import RGBColor
from .utils import handle_args
from pathlib import Path
import asyncio

# # Clear mac terminal
# print('\033c', end=None)


def start_app():
    from_args, is_help_menu = handle_args(sys.argv)
    if is_help_menu:
        return

    job_title_input = input('Job title: ')
    employer_name_input = input('Employer name: ')
    annc_ctrl_num_input = ''
    hr_mng_name_input = ''
    hr_mng_address1_input = ''
    hr_mng_address2_input = ''
    hr_mng_phone_input = ''
    hr_mng_email_input = ''

    if from_args['include_anncmnt_num']:
        annc_ctrl_num_input = input('Announcement/control numbers: ')

    if from_args['include_to_block']:
        hr_mng_name_input = input('HR maneger name: ')
        hr_mng_address1_input = input('HR maneger address, line 1: ')
        hr_mng_address2_input = input('HR maneger address, line 2: ')
        hr_mng_phone_input = input('HR maneger phone number: ')
        hr_mng_email_input = input('HR maneger email address: ')

    user_input = {
        'employer_name': employer_name_input,
        'job_title': job_title_input,
        'annc_ctrl_num': annc_ctrl_num_input,
        'hr_mng_name': hr_mng_name_input,
        'hr_mng_address1': hr_mng_address1_input,
        'hr_mng_address2': hr_mng_address2_input,
        'hr_mng_phone': hr_mng_phone_input,
        'hr_mng_email': hr_mng_email_input
    }

    replace_info(user_input, from_args)


def replace_info(user_input, from_args):
    red_color = RGBColor(0xff, 0x00, 0x00)
    black_color = RGBColor(0x00, 0x00, 0x00)
    start_index = 0
    end_index = 0

    abs_path = os.path.abspath(os.path.dirname(__file__))
    file_path = os.path.join(abs_path, "coverLetterAbdulAliyev.docx")

    doc = docx.Document(file_path)

    for j, paragraph in enumerate(doc.paragraphs):
        # print(j)
        # print(paragraph.text)
        # print('\n')

        # Record start and end indexes of attention block to later remove them
        if 'CURRENT_DATE' in paragraph.text:
            start_index = j
        if 'EMPLOYER_EMAIL' in paragraph.text:
            end_index = j

        for i, run in enumerate(paragraph.runs):
            # print(i)
            # print(run.text)
            # print('\n')

            if from_args['include_to_block']:
                if 'CURRENT_DATE' in run.text:
                    now = datetime.datetime.now()
                    run.text = now.strftime("%m/%d/%Y")
                    run.font.color.rgb = black_color

                if 'HR_MNG_FULL_NAME' in run.text:
                    run.text = user_input['hr_mng_name']
                    run.font.color.rgb = black_color

                if 'EMPLOYER_ADDRESS_1' in run.text and user_input['hr_mng_address1']:
                    run.text = user_input['hr_mng_address1']
                    run.font.color.rgb = black_color

                if 'EMPLOYER_ADDRESS_2' in run.text and user_input['hr_mng_address2']:
                    run.text = user_input['hr_mng_address2']
                    run.font.color.rgb = black_color

                if 'EMPLOYER_PHONE' in run.text and user_input['hr_mng_phone']:
                    run.text = user_input['hr_mng_phone']
                    run.font.color.rgb = black_color

                if 'EMPLOYER_EMAIL' in run.text and user_input['hr_mng_email']:
                    run.text = user_input['hr_mng_email']
                    run.font.color.rgb = black_color

                if 'Dear Hiring Manager:' in run.text:
                    run.text = f"Dear {user_input['hr_mng_name']}:"
                    run.font.color.rgb = black_color

            # Remove Announcement/control number
            if "announcement/control numbers" in run.text:

                if not from_args['include_anncmnt_num']:
                    run.text = ''
                else:
                    run.text = f" (announcement/control numbers: {user_input['annc_ctrl_num']})"
                    run.font.color.rgb = black_color
                    run.font.bold = False

            if 'EMPLOYER_NAME' in run.text:
                run.text = user_input['employer_name']
                run.font.color.rgb = black_color

            if 'JOB_TITLE' in run.text:
                run.text = user_input['job_title']
                run.font.color.rgb = black_color

        # # just for testing to read runs after modifications
        # for i, run in enumerate(paragraph.runs):
        #     print(i)
        #     print(run.text)
        #     print('\n')

    # Remove HR Manager attention block
    if not from_args['include_to_block']:
        for paragraph in doc.paragraphs[start_index:end_index+1]:
            delete_paragraph(paragraph)

    # # Remove Announcement/control number
    # if not from_args['include_to_block']:
    #     for paragraph in doc.paragraphs[start_index:end_index+1]:
    #         delete_paragraph(paragraph)

    home_path = os.environ['HOME']

    # Create job_applications folder under Documents, if not already there
    job_applications_folder_path = os.path.join(
        home_path, "Documents", "job_applications")
    Path(job_applications_folder_path).mkdir(parents=True, exist_ok=True)

    # Create folder with user provided employer anme under job_applications, if not already there
    job_folder_path = os.path.join(
        job_applications_folder_path, user_input['employer_name'])
    Path(job_folder_path).mkdir(parents=True, exist_ok=True)

    new_file_name = "coverLetterAbdulAliyev.docx"

    # Rename a file if it already exists, otherwise it will be overwritten
    for _ in range(100):
        file_exists = os.path.exists(
            os.path.join(job_folder_path, new_file_name))

        if file_exists:
            split_file_name = new_file_name.split(".")
            split_file_name[0] = split_file_name[0] + "_copy"
            new_file_name = ".".join(split_file_name)

    file_path = os.path.join(job_folder_path, new_file_name)

    try:
        doc.save(file_path)
    finally:
        os.system(f"open '{file_path}'")


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
